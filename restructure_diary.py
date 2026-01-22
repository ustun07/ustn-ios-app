#!/usr/bin/env python3
"""
Staj Defteri Yeniden Yapılandırma Script'i
Bu script, mevcut staj defterinden günlük verileri alıp her gün için özgün içerikler üreterek
yeni bir Word belgesi oluşturur.
"""

import re
import random
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def read_original_document(file_path):
    """Orijinal Word belgesini okur ve günleri parse eder."""
    doc = Document(file_path)
    full_text = '\n'.join([para.text for para in doc.paragraphs])
    
    # Günleri ayır
    day_pattern = r'(\d+)\. Gün\s*\n*Yapılan Çalışmanın Konusu\s*:\s*([^\n]+)'
    days = re.findall(day_pattern, full_text)
    
    print(f"Toplam {len(days)} gün bulundu.")
    return days

def generate_unique_work_paragraphs(day_num, topic):
    """Her gün için özgün 3 paragraf oluşturur."""
    
    # Farklı başlangıç şablonları
    intro_templates = [
        "Bugün {topic} üzerinde yoğunlaştım ve kapsamlı çalışmalar gerçekleştirdim.",
        "Günün ana odak noktası {topic} oldu ve bu konuda detaylı araştırmalar yaptım.",
        "{topic} ile ilgili teorik bilgileri pratik uygulamalarla pekiştirdim.",
        "Sabah saatlerinde {topic} konusuna giriş yaptım ve temel kavramları öğrendim.",
        "{topic} alanında yeni beceriler kazandım ve mevcut bilgilerimi derinleştirdim.",
        "Bugünkü çalışmalarımı {topic} odağında yürüttüm ve önemli ilerlemeler kaydettim.",
        "{topic} konusunda hem teorik hem de pratik çalışmalar yaparak deneyim kazandım.",
        "Gün boyunca {topic} üzerine odaklandım ve farklı yaklaşımları deneme fırsatı buldum.",
    ]
    
    # Orta paragraf şablonları
    middle_templates = [
        "Proje gereksinimlerini analiz ederek, hangi adımların öncelikli olduğunu belirledim. Geliştirme sürecinde karşılaşabileceğim potansiyel zorlukları öngörmeye çalıştım ve bunlara yönelik çözüm stratejileri geliştirdim. Ekip üyeleriyle düzenli iletişim kurarak, projenin genel gidişatı hakkında bilgi alışverişinde bulundum.",
        "İlgili dokümantasyonları inceleyerek konuyu derinlemesine anlamaya çalıştım. Öğrendiğim yeni teknikleri küçük test senaryolarında deneyerek pratik yapma imkanı buldum. Farklı yaklaşımların avantaj ve dezavantajlarını karşılaştırarak en uygun yöntemi belirlemeye çalıştım.",
        "Sistemin mevcut yapısını analiz ederek iyileştirme fırsatlarını tespit ettim. Kod kalitesini artırmak için best practice'leri araştırdım ve uygulamaya çalıştım. Performans optimizasyonu açısından kritik noktaları belirleyip notlar aldım.",
        "Kullanıcı deneyimini geliştirmek için farklı tasarım alternatifleri üzerinde düşündüm. Arayüz elemanlarının erişilebilirliğini kontrol ettim ve gerekli düzenlemeleri planladım. Modern tasarım prensiplerine uygun çözümler araştırarak fikir edindim.",
        "Veri akışını optimize etmek için farklı mimariler inceledim. Ölçeklenebilirlik açısından en uygun yapıyı belirlemek için karşılaştırmalar yaptım. Güvenlik standartlarına uygunluğu kontrol ederek notlar çıkardım.",
        "Entegrasyon süreçlerini detaylı şekilde inceleyerek öğrendim. API dokümantasyonlarını okuyarak farklı servislerin nasıl çalıştığını anladım. Test ortamında çeşitli senaryoları deneyerek sistemin davranışını gözlemledim.",
        "Kod tabanını refactor etmek için fırsatlar aradım ve değerlendirdim. Tekrarlayan kodları tespit ederek modüler yapıya geçiş planları yaptım. Clean code prensiplerini uygulayarak okunabilirliği artırmaya çalıştım.",
        "Hata ayıklama tekniklerini uygulayarak sistemdeki potansiyel sorunları araştırdım. Loglama mekanizmalarını inceleyerek sorun takibinin nasıl yapıldığını öğrendim. Debugging sürecinde kullanılabilecek araçları keşfettim.",
        "Veritabanı şemasını inceleyerek veri modelini anlamaya çalıştım. Query optimizasyonu için farklı yaklaşımları araştırdım. İndeksleme stratejilerini öğrenerek performans iyileştirmeleri hakkında fikir edindim.",
        "Kullanıcı geri bildirimlerini değerlendirerek önceliklendirme yaptım. Feature isteklerini analiz ederek teknik gereksinimleri çıkarmaya çalıştım. Ürün geliştirme sürecinde kullanıcı odaklı düşünmenin önemini kavradım.",
    ]
    
    # Son paragraf şablonları
    closing_templates = [
        "Gün sonunda elde ettiğim bilgileri notlara döktüm ve öğrendiklerimi pekiştirdim. Yarın için çalışma planını belirleyerek hangi konulara odaklanacağımı netleştirdim. Bu sayede sürekli gelişen bir öğrenme süreci oluşturduğumu gözlemledim.",
        "Edindiğim deneyimleri bir sonraki aşamada nasıl kullanabileceğimi düşündüm. Yapılacaklar listesini güncelleyerek önceliklendirme yaptım. Takip eden günler için daha verimli bir çalışma düzeni kurgulamaya çalıştım.",
        "Bugünkü kazanımlarımı değerlendirerek güçlü ve zayıf yönlerimi belirledim. İyileştirilmesi gereken alanları tespit edip bunlara odaklanmayı planladım. Sürekli öğrenme ve gelişim için motivasyonumu artırdım.",
        "Öğrendiğim kavramları gerçek dünya senaryolarıyla ilişkilendirerek içselleştirdim. Teorik bilgilerin pratik uygulamalardaki karşılığını anlamlandırdım. Bu bütünsel bakış açısı sayesinde konuyu daha iyi kavradım.",
        "Elde ettiğim bilgileri dokümante ederek ileride referans olması için kayıt altına aldım. Karşılaştığım zorlukları ve çözüm yollarını not ettim. Bu sayede benzer durumlarla karşılaştığımda daha hazırlıklı olacağımı düşünüyorum.",
        "Günün sonunda tamamladığım görevleri gözden geçirerek ilerleme kaydettim. Açık kalan konuları belirlep bir sonraki gün için hedef oluşturdum. Böylece sürekli ilerleyen bir çalışma ritmi yakaladım.",
        "Öğrendiklerimi ekip arkadaşlarımla paylaşarak bilgi alışverişinde bulundum. Farklı bakış açılarını dinleyerek perspektifimi genişlettim. İşbirlikçi çalışma ortamının öğrenmeye katkısını deneyimledim.",
        "Bugünkü çalışmaların genel proje içindeki yerini değerlendirdim. Her adımın büyük resme nasıl katkı sağladığını anlamlandırdım. Bu bütünsel görüş sayesinde motivasyonumu korumayı başardım.",
        "Karşılaştığım yeni terimleri ve kavramları öğrenerek kelime dağarcığımı geliştirdim. Teknik dokümantasyon okuma becerilerimi ilerletim. Profesyonel gelişimime katkı sağlayan bir gün geçirdim.",
        "Gün içinde aldığım küçük notları derleyerek kapsamlı bir özet çıkardım. Problem çözme yaklaşımlarımı analiz ederek daha etkili yöntemler geliştirmeye çalıştım. Refleksif düşünme sayesinde farkındalığımı artırdım.",
    ]
    
    # Rastgele seçim yap ama belirli bir seed ile her gün için tutarlılık sağla
    random.seed(day_num * 7)
    intro = random.choice(intro_templates).format(topic=topic.lower())
    
    random.seed(day_num * 13)
    middle = random.choice(middle_templates)
    
    random.seed(day_num * 19)
    closing = random.choice(closing_templates)
    
    return [intro, middle, closing]

def generate_problem_section(day_num):
    """Her gün için özgün problem bölümü oluşturur."""
    
    problem_templates = [
        "Bugün belirgin bir teknik sorunla karşılaşmadım ancak bazı kavramları anlamakta zorlandım. Dokümantasyon eksikliği nedeniyle araştırma süreci biraz uzadı. İnternet kaynaklarından ve ekip üyelerinden destek alarak sorunu aştım.",
        "Geliştirme ortamında konfigürasyon hatası nedeniyle zaman kaybı yaşadım. Hata mesajlarını analiz ederek kaynağı tespit ettim ve gerekli düzeltmeleri yaptım. Bu süreç troubleshooting becerilerimi geliştirmeme yardımcı oldu.",
        "Yeni bir teknoloji öğrenirken başlangıçta kavramları ilişkilendirmekte güçlük çektim. Adım adım ilerleyerek ve pratik örnekler yaparak konuyu anlamayı başardım. Öğrenme eğrisinin doğal bir parçası olduğunu fark ettim.",
        "Karmaşık bir yapıyı anlamak için beklenenden fazla zaman harcadım. Sistemin farklı bileşenlerinin nasıl etkileştiğini çözmek zaman aldı. Görselleştirme ve şematik çizimler yaparak durumu netleştirdim.",
        "Versiyon uyumsuzluğu nedeniyle beklenmeyen hatalarla karşılaştım. Bağımlılık yönetiminin önemini bir kez daha deneyimledim. Doğru versiyonları kullanarak sorunu çözdüm.",
        "Test sürecinde beklenmeyen bir davranış gözlemledim. Root cause analysis yaparak hatanın kaynağını buldum. Debug araçlarını etkili kullanarak sorunu çözdüm.",
        "Performans sorunları tespit ettim ve optimizasyon gerektiren alanları belirledim. Profiling araçları kullanarak darboğazları tespit ettim. İyileştirme stratejileri geliştirdim.",
        "API entegrasyonu sırasında authentication hatası aldım. Dokümantasyonu tekrar inceleyerek doğru parametreleri kullandım. Token yönetimi konusunda deneyim kazandım.",
        "Kod review sürecinde belirlenen iyileştirme önerileri üzerinde çalıştım. Best practice'lere uygunluğu artırmak için düzenlemeler yaptım. Code quality konusunda yeni bakış açıları kazandım.",
        "Kompleks bir algoritmanın implementasyonu sırasında mantıksal hata yaptım. Adım adım debug ederek hatayı buldum ve düzelttim. Algoritma tasarımında daha dikkatli olmam gerektiğini öğrendim.",
    ]
    
    random.seed(day_num * 23)
    return random.choice(problem_templates)

def generate_evaluation(day_num, topic):
    """Her gün için özgün değerlendirme oluşturur."""
    
    evaluation_templates = [
        "Bugün {topic} alanında önemli kazanımlar elde ettim. Teorik bilgileri pratiğe dökmek, kavramları daha iyi anlamama yardımcı oldu. İlerleyen günlerde bu bilgileri geliştirerek kullanacağım.",
        "{topic} üzerine yaptığım çalışmalar, profesyonel gelişimime katkı sağladı. Farklı yaklaşımları deneme fırsatı buldum ve deneyim kazandım. Bu deneyimlerin gelecekte faydalı olacağına inanıyorum.",
        "Günün sonunda {topic} konusunda kendimi daha yetkin hissediyorum. Karşılaştığım zorluklar problem çözme becerilerimi geliştirdi. Sürekli öğrenmenin önemini bir kez daha deneyimledim.",
        "{topic} ile ilgili çalışmalar sayesinde teknik becerilerimi artırdım. Pratik uygulama yapmanın teorik bilgiden çok daha etkili olduğunu gördüm. Motivasyonumu yüksek tutarak ilerliyorum.",
        "Bugün {topic} alanında edindiğim bilgiler temel oluşturdu. İleriki aşamalarda bu temeller üzerine inşa edeceğim. Sabırlı ve sistematik çalışmanın önemini kavradım.",
        "{topic} konusunda derinlemesine çalışma fırsatı buldum. Her aşamada yeni şeyler öğrenmek beni heyecanlandırıyor. Bu pozitif enerjiyi koruyarak devam edeceğim.",
        "Gün boyunca {topic} üzerinde yoğunlaştım ve verimli bir süreç geçirdim. Öğrenme hedeflerime ulaşmak için doğru yolda olduğumu düşünüyorum. Kendi gelişimimi gözlemlemek motivasyon kaynağım.",
        "{topic} alanındaki çalışmalarım, beklentilerimin üzerinde sonuçlar verdi. Planlı ve düzenli çalışmanın getirilerini gözlemledim. Staj sürecinden maksimum fayda sağladığımı düşünüyorum.",
        "Bugün {topic} konusunda pratik deneyim kazanmak çok değerliydi. Gerçek dünya problemleriyle uğraşmak bana farklı bir perspektif kazandırdı. Bu deneyimlerin kariyer gelişimime katkı sağlayacağından eminim.",
        "{topic} üzerine çalışırken hem zorlandım hem de keyif aldım. Öğrenme sürecinin dinamik olması beni daha aktif kılıyor. Her gün yeni bir şeyler öğrendiğim için minnettarım.",
    ]
    
    random.seed(day_num * 29)
    return random.choice(evaluation_templates).format(topic=topic.lower())

def create_new_document(days_data):
    """Yeni Word belgesi oluşturur."""
    doc = Document()
    
    # Stil ayarları
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    for day_num, topic in days_data:
        day_num = int(day_num)
        
        # Gün başlığı
        heading = doc.add_paragraph(f'{day_num}. Gün')
        heading.style = 'Heading 1'
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Yapılan Çalışmanın Konusu
        topic_para = doc.add_paragraph()
        topic_para.add_run('Yapılan Çalışmanın Konusu : ').bold = True
        topic_para.add_run(topic)
        
        # Yapılan Çalışmalar
        work_heading = doc.add_paragraph()
        work_heading.add_run('Yapılan Çalışmalar:').bold = True
        
        # 3 özgün paragraf ekle
        work_paragraphs = generate_unique_work_paragraphs(day_num, topic)
        for para_text in work_paragraphs:
            doc.add_paragraph(para_text)
        
        # Karşılaşılan Problemler
        problem_heading = doc.add_paragraph()
        problem_heading.add_run('Karşılaşılan Problemler / Çözümler:').bold = True
        
        problem_text = generate_problem_section(day_num)
        doc.add_paragraph(problem_text)
        
        # Günlük Değerlendirme
        eval_heading = doc.add_paragraph()
        eval_heading.add_run('Günlük Değerlendirme:').bold = True
        
        eval_text = generate_evaluation(day_num, topic)
        doc.add_paragraph(eval_text)
        
        # Gün sonu bilgileri
        doc.add_paragraph(f'\nTarih : ____/____/20____      Kaşe / İmza      Sayfa No : {day_num}')
        
        # Sayfa sonu (son gün hariç)
        if day_num < len(days_data):
            doc.add_page_break()
        
        print(f"{day_num}. gün işlendi")
    
    return doc

def main():
    """Ana fonksiyon."""
    input_file = '/Users/selimcanustun/Downloads/staj_defteri_76_gun_final.docx'
    output_file = '/Users/selimcanustun/Downloads/staj_defteri_76_gun_yeniden_yapilandirilmis.docx'
    
    print("Orijinal dosya okunuyor...")
    days_data = read_original_document(input_file)
    
    if not days_data:
        print("HATA: Günler bulunamadı!")
        return
    
    print(f"\nYeni belge oluşturuluyor ({len(days_data)} gün)...")
    new_doc = create_new_document(days_data)
    
    print(f"\nKaydediliyor: {output_file}")
    new_doc.save(output_file)
    
    print("\n✅ İşlem tamamlandı!")
    print(f"Yeni dosya: {output_file}")

if __name__ == '__main__':
    main()
